#!/usr/bin/env python3
"""Extract DOCX paragraphs and tables with stable indexes and SHA-256 hashes."""
import argparse, hashlib, json
from docx import Document

def digest(text):
    return hashlib.sha256(text.encode('utf-8')).hexdigest()

def extract(path):
    doc = Document(path)
    items = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip(): items.append({'kind':'paragraph','index':i,'text':p.text,'sha256':digest(p.text)})
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            text = ' | '.join(cell.text for cell in row.cells)
            items.append({'kind':'table_row','table':ti,'index':ri,'text':text,'sha256':digest(text)})
    return items

if __name__ == '__main__':
    ap = argparse.ArgumentParser(); ap.add_argument('input'); ap.add_argument('-o','--output')
    args = ap.parse_args(); data = {'input':args.input,'items':extract(args.input)}
    out = json.dumps(data, ensure_ascii=False, indent=2)
    if args.output: open(args.output,'w',encoding='utf-8').write(out+'\n')
    else: print(out)
