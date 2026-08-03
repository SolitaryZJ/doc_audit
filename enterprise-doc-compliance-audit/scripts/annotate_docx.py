#!/usr/bin/env python3
"""Create a highlighted DOCX copy from simple paragraph-index findings."""
import argparse, json, shutil
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def annotate(src, dst, findings):
    shutil.copy2(src, dst); doc = Document(dst); applied=[]
    for f in findings:
        loc=f.get('location',{}); idx=loc.get('paragraph_index');
        if not isinstance(idx,int) or idx < 0 or idx >= len(doc.paragraphs): continue
        p=doc.paragraphs[idx]; needle=loc.get('text') or f.get('excerpt','')
        if needle and needle not in p.text: continue
        for run in p.runs: run.font.highlight_color=WD_COLOR_INDEX.YELLOW
        applied.append(f.get('finding_id'))
    doc.save(dst); return applied

if __name__ == '__main__':
    ap=argparse.ArgumentParser(); ap.add_argument('input'); ap.add_argument('output'); ap.add_argument('findings_json')
    a=ap.parse_args(); findings=json.load(open(a.findings_json,encoding='utf-8'))
    print(json.dumps({'annotated':annotate(a.input,a.output,findings)},ensure_ascii=False))
